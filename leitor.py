import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
import time
import io
import datetime
from docx import Document
from docx.shared import Pt

# Configuração do Tesseract (Comentado para rodar na nuvem)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

ESTADOS_BR = {
    "ACRE": "AC", "ALAGOAS": "AL", "AMAPÁ": "AP", "AMAZONAS": "AM", "BAHIA": "BA",
    "CEARÁ": "CE", "DISTRITO FEDERAL": "DF", "ESPÍRITO SANTO": "ES", "GOIÁS": "GO",
    "MARANHÃO": "MA", "MATO GROSSO": "MT", "MATO GROSSO DO SUL": "MS", "MINAS GERAIS": "MG",
    "PARÁ": "PA", "PARAÍBA": "PB", "PARANÁ": "PR", "PERNAMBUCO": "PE", "PIAUÍ": "PI",
    "RIO DE JANEIRO": "RJ", "RIO GRANDE DO NORTE": "RN", "RIO GRANDE DO SUL": "RS",
    "RONDÔNIA": "RO", "RORAIMA": "RR", "SANTA CATARINA": "SC", "SÃO PAULO": "SP",
    "SERGIPE": "SE", "TOCANTINS": "TO"
}

# ==========================================
# Interface do App - BARRA LATERAL
# ==========================================
st.set_page_config(page_title="Gerador de Iniciais", layout="wide")
st.sidebar.title("🔐 Configurações")

usuario_salvo = st.secrets.get("COBRARE_USER", "")
senha_salva = st.secrets.get("COBRARE_PASSWORD", "")

usuario_cobrare = st.sidebar.text_input("Usuário do COBRARE", value=usuario_salvo)
senha_cobrare = st.sidebar.text_input("Senha do COBRARE", type="password", value=senha_salva)

# ==========================================
# Módulos de Extração e Automação Web
# ==========================================
def extrair_texto_hibrido(arquivo_bytes):
    texto_completo = ""
    try:
        doc = fitz.open(stream=arquivo_bytes, filetype="pdf")
        for num_pagina in range(len(doc)):
            pagina = doc.load_page(num_pagina)
            texto_nativo = pagina.get_text()
            
            # Se não tiver texto, é um PDF escaneado. Vamos "tirar uma foto" da página.
            if len(texto_nativo.strip()) < 50:
                # O matrix 2x2 dobra a resolução da imagem para o robô ler com perfeição
                pix = pagina.get_pixmap(matrix=fitz.Matrix(2, 2)) 
                img_data = pix.tobytes("png")
                imagem = Image.open(io.BytesIO(img_data))
                
                texto_ocr = pytesseract.image_to_string(imagem, lang='por')
                texto_completo += texto_ocr + "\n"
            else:
                texto_completo += texto_nativo + "\n"
        doc.close()
    except Exception as e:
        return f"Erro ao processar o arquivo: {e}"
    return texto_completo

def minerar_dados_confissao(texto_bruto):
    dados = {"credor": "Não encontrado", "devedor": "Não encontrado", "cpf_devedor": None, "cidade_comarca": "Não encontrada"}
    texto_limpo = texto_bruto.replace('\n', ' ').replace('  ', ' ')
    
    # Busca o Credor
    busca_credor = re.search(r"de um lado,\s*(.*?),\s*(?:representada neste ato|doravante denominada 1ª)", texto_limpo, re.IGNORECASE)
    if busca_credor: 
        dados["credor"] = busca_credor.group(1).strip()
        
    # Busca o Devedor (Agora a vírgula antes de "doravante" é opcional)
    busca_devedor = re.search(r"de outro lado,\s*(.*?)(?:,|\s)*doravante denominad[oa]", texto_limpo, re.IGNORECASE)
    if busca_devedor:
        dados["devedor"] = busca_devedor.group(1).strip()
        
        # Busca o CPF (Agora o "sob o" é opcional e aceita diferentes formatos de "nº")
        busca_cpf = re.search(r"CPF(?:/MF)?\s*(?:sob o\s*)?n[º°o]?\s*([\d\.\-]+)", dados["devedor"], re.IGNORECASE)
        if busca_cpf: 
            dados["cpf_devedor"] = busca_cpf.group(1).strip()
        
        # Busca a Cidade (Pega o texto antes do " - RS/SP/etc", ignorando números de CEP que possam estar grudados)
        busca_cidade = re.search(r"([A-Za-zÀ-ÿ\s]+)\s*-\s*[A-Z]{2}", dados["devedor"])
        if busca_cidade:
            cidade_suja = busca_cidade.group(1).strip()
            # Corta pela última vírgula para isolar apenas o nome da cidade
            cidade_limpa = cidade_suja.split(',')[-1].strip()
            # Remove qualquer número (como CEP) que tenha vindo junto
            cidade_limpa = re.sub(r'\d+', '', cidade_limpa).strip()
            dados["cidade_comarca"] = cidade_limpa.upper()
            
    return dados

def buscar_endereco_cobrare(pesquisa_devedor):
    chrome_options = Options()
    chrome_options.add_argument("--headless=new") 
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920,1080") 
    chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36")
    
    driver = webdriver.Chrome(options=chrome_options)
    wait = WebDriverWait(driver, 15)
    
    try:
        driver.get("https://cobrare.atenta.pro/controleDivida/menu")
        wait.until(EC.presence_of_element_located((By.ID, "login"))).send_keys(usuario_cobrare)
        driver.find_element(By.ID, "password").send_keys(senha_cobrare)
        driver.find_element(By.ID, "password").submit()
        
        wait.until(EC.element_to_be_clickable((By.XPATH, "//a[@data-toggle='dropdown' and contains(., 'Movimentação')]"))).click()
        time.sleep(1)
        wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(@href, '/controleDivida/negociacao')]"))).click()
        
        campo_busca = wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@aria-controls='negociacao']")))
        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", campo_busca)
        time.sleep(1)
        
        campo_busca.click()
        campo_busca.clear()
        
        cpf_limpo = pesquisa_devedor.replace(".", "").replace("-", "")
        for numero in cpf_limpo:
            campo_busca.send_keys(numero)
            time.sleep(0.1)
            
        campo_busca.send_keys(Keys.ENTER)
        time.sleep(5)
        
        botoes_editar = driver.find_elements(By.XPATH, "//a[contains(@href, '/edit') and contains(@class, 'btn-mini')]")
        if len(botoes_editar) > 0:
            botoes_editar[0].click()
        else:
            raise Exception("Devedor não encontrado na tabela.")
        
        aba_info = wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(@id, 'ui-id-') and contains(@class, 'ui-tabs-anchor')]")))
        aba_info.click()
        wait.until(EC.presence_of_element_located((By.ID, "txEndereco")))
        
        cidade_select = Select(driver.find_element(By.ID, "cidadeCombo"))
        estado_select = Select(driver.find_element(By.ID, "estadoCombo"))
        
        estado_completo = estado_select.first_selected_option.get_attribute("textContent").strip().upper()
        
        dados_sistema = {
            "latitude": driver.find_element(By.ID, "nrLatitude").get_attribute("value"),
            "longitude": driver.find_element(By.ID, "nrLongitude").get_attribute("value"),
            "logradouro": driver.find_element(By.ID, "txEndereco").get_attribute("value"),
            "numero": driver.find_element(By.ID, "nrEndereco").get_attribute("value"),
            "complemento": driver.find_element(By.ID, "dsComplemento").get_attribute("value"),
            "bairro": driver.find_element(By.ID, "nmBairro").get_attribute("value"),
            "cep": driver.find_element(By.ID, "nrCep").get_attribute("value"),
            "cidade": cidade_select.first_selected_option.get_attribute("textContent").strip(),
            "estado": ESTADOS_BR.get(estado_completo, estado_completo) 
        }
        return dados_sistema
    except Exception as e:
        return str(e)
    finally:
        driver.quit()

def atualizar_endereco_devedor(texto_devedor, endereco_cobrare, cidade_comarca):
    if not isinstance(endereco_cobrare, dict): return texto_devedor
    
    comp = f" - {endereco_cobrare['complemento']}" if endereco_cobrare['complemento'] else ""
    cidade_estado = f"{endereco_cobrare['cidade']} - {endereco_cobrare['estado']}"
    lat, lon = endereco_cobrare.get('latitude', '').strip(), endereco_cobrare.get('longitude', '').strip()
    coords = f", Latitude {lat}, Longitude {lon}" if lat and lon else ""
    
    novo_endereco = f"{endereco_cobrare['logradouro']}, nº {endereco_cobrare['numero']}{comp}, Bairro {endereco_cobrare['bairro']}, CEP {endereco_cobrare['cep']}{coords}, {cidade_estado}"
    
    # Busca inteligente: acha a palavra residente ou domiciliado e corta o texto logo em seguida!
    match = re.search(r"(residente|domiciliado)(.*?)(em|na|no)\s+", texto_devedor, re.IGNORECASE)
    if match:
        idx = match.end()
        # Cola o texto original até a palavra "na/no/em" e adiciona o endereço novo do COBRARE
        return texto_devedor[:idx] + novo_endereco
    else:
        # Se não achar nenhuma das palavras (fallback de segurança)
        return texto_devedor + f", residente e domiciliado em {novo_endereco}"

def minerar_dados_detran(texto_bruto):
    dados = {"marca": "Não encontrada", "placa": "Não encontrada", "ano_modelo": "Não encontrado", "cor": "Não encontrada", "renavam": "Não encontrado", "chassi": "Não encontrado"}
    buscas = {
        "placa": r"Placa:\s*([A-Z0-9]{7})", "chassi": r"Chassi:\s*([A-Z0-9]+)", "marca": r"Marca:\s*(.+)",
        "renavam": r"RENAVAM:\s*([0-9]+)", "ano_modelo": r"Fabricação/Modelo:\s*([\d]{4}\s*/\s*[\d]{4})", "cor": r"Cor:\s*([A-Za-zÀ-ÿ]+)"
    }
    for chave, padrao in buscas.items():
        match = re.search(padrao, texto_bruto, re.IGNORECASE)
        if match:
            valor = match.group(1).strip()
            if chave == "ano_modelo": valor = valor.replace(" ", "")
            if chave == "cor": valor = valor.lower()
            dados[chave] = valor
    return dados

def minerar_dados_imovel(texto_bruto):
    dados = {"cartorio": "Não encontrado", "matricula": "Não encontrada"}
    
    # Limpa a leitura OCR, removendo sujeiras visuais
    texto_limpo = texto_bruto.replace('$', '').replace('^{\\circ}', 'º').replace('{', '').replace('}', '').replace('\\', '')
    
    # Busca o Cartório: Aceita "Cartório" com ou sem acento
    busca_cartorio = re.search(r"(\d+(?:º|°|o)\s*Cart[oó]rio\s*-\s*[A-Za-zÀ-ÿ]+)", texto_limpo, re.IGNORECASE)
    if busca_cartorio:
        dados["cartorio"] = busca_cartorio.group(1).strip()
        
    # Busca a Matrícula: Acha o CPF (11+ números), ignora espaços e pega a matrícula (3 a 8 números)
    busca_matricula = re.search(r"\d{11,14}\s+(\d{3,8})\s*(?:Sim|Não|S1m|Nao)?", texto_limpo, re.IGNORECASE)
    if busca_matricula:
        dados["matricula"] = busca_matricula.group(1).strip()
        
    return dados

# ==========================================
# Módulo 4: Preenchimento Avançado do Word
# ==========================================
def aplicar_estilo_garamond(run):
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return run

def forcar_paragrafo_bold(paragrafo):
    texto_completo = paragrafo.text
    paragrafo.clear()
    run = aplicar_estilo_garamond(paragrafo.add_run(texto_completo))
    run.bold = True

def formatar_qualificacao(paragrafo, texto_qualificacao):
    match = re.search(r',\s*(inscrit|brasileir|pessoa jurídica|com sede|residente|portador|solteir|casad|divorciad|viúv|agricultor|empresári|menor)', texto_qualificacao, re.IGNORECASE)
    
    if match:
        idx = match.start()
        nome = texto_qualificacao[:idx].strip().upper()
        resto = texto_qualificacao[idx:]
    else:
        partes = texto_qualificacao.split(',', 1)
        nome = partes[0].strip().upper()
        resto = "," + partes[1] if len(partes) > 1 else ""

    if '[QUALIFICAÇÃO COMPLETA]' in paragrafo.text:
        texto_antes, texto_depois = paragrafo.text.split('[QUALIFICAÇÃO COMPLETA]', 1)
        paragrafo.clear()

        if texto_antes: aplicar_estilo_garamond(paragrafo.add_run(texto_antes))
        
        run_nome = aplicar_estilo_garamond(paragrafo.add_run(nome))
        run_nome.bold = True
        
        if resto: aplicar_estilo_garamond(paragrafo.add_run(resto))
        if texto_depois: aplicar_estilo_garamond(paragrafo.add_run(texto_depois))

def gerar_documento_word(caminho_modelo, comarca, credor, devedor_qualificado, lista_veiculos, lista_imoveis):
    doc = Document(caminho_modelo)
    qualificacao_count = 0
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    hoje = datetime.date.today()
    data_formatada = f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"
    
    for p in doc.paragraphs:
        texto_upper = p.text.strip().upper()
        
        if '[COMARCA]' in p.text:
            p.text = p.text.replace('[COMARCA]', comarca.upper())
            forcar_paragrafo_bold(p)
            continue
            
        if '[QUALIFICAÇÃO COMPLETA]' in p.text:
            if qualificacao_count == 0:
                formatar_qualificacao(p, credor)
                qualificacao_count += 1
            else:
                formatar_qualificacao(p, devedor_qualificado)
            continue
                
        if 'Santa Cruz do Sul/RS,' in p.text:
            nova_data = re.sub(r'Santa Cruz do Sul/RS,.*', f'Santa Cruz do Sul/RS, {data_formatada}.', p.text)
            p.clear()
            aplicar_estilo_garamond(p.add_run(nova_data))
            continue
                
        # Formatação Unificada: Veículos e/ou Imóveis
        if 'd) A expedição de' in p.text and (lista_veiculos or lista_imoveis):
            p.clear() 
            aplicar_estilo_garamond(p.add_run("d) Para a efetivação da penhora, a Exequente indica "))
            
            # Adiciona os imóveis (se houver)
            if lista_imoveis:
                textos_imoveis = [f"matrícula nº {im['matricula']}, registrado no Registro de Imóveis de {im['cartorio']}" for im in lista_imoveis]
                if len(textos_imoveis) == 1:
                    texto_imovel_completo = "o imóvel de " + textos_imoveis[0]
                else:
                    texto_imovel_completo = "os imóveis de " + " e ".join(textos_imoveis)
                
                aplicar_estilo_garamond(p.add_run(texto_imovel_completo))
                
                # Se tiver veículo também, faz a ponte
                if lista_veiculos:
                    aplicar_estilo_garamond(p.add_run(", bem como os seguintes veículos de propriedade do Executado: "))
                else:
                    aplicar_estilo_garamond(p.add_run(" de propriedade do Executado."))
            
            # Adiciona os veículos (se houver)
            elif lista_veiculos:
                aplicar_estilo_garamond(p.add_run("os seguintes veículos de propriedade do Executado: "))
            
            if lista_veiculos:
                romanos = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
                for i, v in enumerate(lista_veiculos):
                    idx_romano = romanos[i] if i < len(romanos) else str(i + 1)
                    
                    run_bold = aplicar_estilo_garamond(p.add_run(f"{idx_romano} - {v['marca']}"))
                    run_bold.bold = True
                    
                    separator = "; " if i < len(lista_veiculos) - 1 else "."
                    aplicar_estilo_garamond(p.add_run(f", de placa {v['placa']}, ano/modelo {v['ano_modelo']}, cor {v['cor']}, possui o RENAVAM {v['renavam']} e o chassi {v['chassi']}{separator} "))
            continue
            
        if texto_upper.startswith("I – DOS FATOS") or \
           texto_upper.startswith("II – DO DIREITO") or \
           texto_upper.startswith("III – DOS PEDIDOS") or \
           texto_upper.startswith("DIANTE DO EXPOSTO") or \
           texto_upper.startswith("AÇÃO DE EXECUÇÃO POR QUANTIA CERTA") or \
           (texto_upper == "CLEIDIMARA DA SILVA FLORES") or \
           (texto_upper.startswith("OAB/RS")):
            forcar_paragrafo_bold(p)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# Placeholder para Upload no OneDrive
# ==========================================
def enviar_para_onedrive(credor, devedor, arquivos):
    nome_pasta = f"{credor.upper()} X {devedor.upper()}"
    pass

# ==========================================
# Execução Principal
# ==========================================
st.title("📄 Gerador de Iniciais Automatizado")

st.subheader("1. Documentos Obrigatórios")
confissao_file = st.file_uploader("Anexar Confissão de Dívida (PDF)", type=["pdf"])

st.subheader("2. Documentos Opcionais")
detran_files = st.file_uploader("Anexar Certidão(ões) do Detran (PDF) - Opcional", type=["pdf"], accept_multiple_files=True)
imoveis_files = st.file_uploader("Anexar Certidão(ões) de Imóveis (PDF) - Opcional", type=["pdf"], accept_multiple_files=True)
documentos_extras = st.file_uploader("Anexar Documentos Extras (Qualquer formato) - Opcional", accept_multiple_files=True)

if st.button("Processar e Gerar Inicial"):
    if confissao_file:
        st.success("Iniciando montagem da inicial...")
        
        with st.spinner("Minerando Confissão de Dívida..."):
            texto_confissao = extrair_texto_hibrido(confissao_file.getvalue())


            

            dados_minerados = minerar_dados_confissao(texto_confissao)
            qualificacao_devedor = dados_minerados["devedor"]

        
            
        if dados_minerados['cpf_devedor']:
            with st.spinner("Buscando endereço no COBRARE..."):
                endereco_cobrare = buscar_endereco_cobrare(dados_minerados['cpf_devedor'])
                if isinstance(endereco_cobrare, str):
                    st.warning(f"Não foi possível buscar no COBRARE. Usando endereço do contrato. (Erro: {endereco_cobrare})")
                qualificacao_devedor = atualizar_endereco_devedor(qualificacao_devedor, endereco_cobrare, dados_minerados['cidade_comarca'])

        lista_veiculos = []
        if detran_files:
            for i, certidao in enumerate(detran_files):
                with st.spinner(f"Processando Certidão Detran {i+1} ({certidao.name})..."):
                    texto_detran = extrair_texto_hibrido(certidao.getvalue())
                    lista_veiculos.append(minerar_dados_detran(texto_detran))
                    
        lista_imoveis = []
        if imoveis_files:
            for i, certidao in enumerate(imoveis_files):
                with st.spinner(f"Processando Certidão Imóvel {i+1} ({certidao.name})..."):
                    texto_imovel = extrair_texto_hibrido(certidao.getvalue())

                    lista_imoveis.append(minerar_dados_imovel(texto_imovel))
                
        with st.spinner("Montando o documento final e formatando estilos..."):
            caminho_modelo_word = "1. Modelo inicial execução (confissão de dívida).docx"
            
            doc_final_bytes = gerar_documento_word(
                caminho_modelo_word,
                dados_minerados['cidade_comarca'],
                dados_minerados['credor'],
                qualificacao_devedor,
                lista_veiculos,
                lista_imoveis
            )
            
            todos_os_arquivos = [confissao_file] + (detran_files if detran_files else []) + (imoveis_files if imoveis_files else []) + (documentos_extras if documentos_extras else [])
            enviar_para_onedrive(dados_minerados['credor'], dados_minerados['devedor'], todos_os_arquivos)

            st.success("✅ Petição Inicial gerada com sucesso!")
            st.download_button(
                label="⬇️ Baixar Inicial Pronta (.docx)",
                data=doc_final_bytes,
                file_name=f"Inicial_Execucao_{dados_minerados['cpf_devedor']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("⚠️ O arquivo da Confissão de Dívida é obrigatório.")